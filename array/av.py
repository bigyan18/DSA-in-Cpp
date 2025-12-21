import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def create_document():
    doc = Document()

    # --- Helper Function to Add Images ---
    def add_image_from_url(document, url, caption, width=4.0):
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                # Add image centered
                image_stream = BytesIO(response.content)
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(width))
                
                # Add caption
                caption_para = document.add_paragraph(caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.runs[0]
                caption_run.italic = True
                caption_run.font.size = Pt(9)
            else:
                document.add_paragraph(f"[Image: {caption} - Could not be downloaded]")
        except Exception as e:
            print(f"Error downloading {caption}: {e}")
            document.add_paragraph(f"[Image: {caption} - Error downloading]")

    # --- Cover Page ---
    # University Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TIANJIN CHENGJIAN UNIVERSITY")
    run.bold = True
    run.font.size = Pt(16)
    
    p = doc.add_paragraph() # Spacer
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("天津城建大学")
    run.bold = True
    run.font.size = Pt(18)
    
    doc.add_paragraph("\n" * 3) # Spacing

    # Assignment Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ASSIGNMENT REPORT")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph("\n" * 2)

    # Student Details (Using a table for alignment)
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    details = [
        ("Student Name:", "Aviyan Thapa"),
        ("Student ID:", "2024GJ0201016"),
        ("Program:", "Civil Engineering"),
        ("Course:", "Art Appreciation"),
        ("Supervisor:", "Jei Wei"),
        ("Date:", "2025.12.22")
    ]
    
    for row, (label, value) in zip(table.rows, details):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        # Adjust column widths if needed, but auto works mostly

    doc.add_page_break()

    # --- Page 2: Aims ---
    doc.add_heading("The Aims of the Project", level=1)
    
    doc.add_paragraph("I would like to express my sincere gratitude to my course instructor for their guidance, encouragement, and valuable feedback throughout the completion of this project. Their support played an important role in helping me understand the subject more deeply and complete this work successfully.")
    
    doc.add_paragraph("The main aim of this project is to study the Song Dynasty (960-1279 CE) as a significant period in the history of Chinese art, with a special focus on landscape painting (shan shui). Through this project, I aim to explore how philosophy, nature, and personal expression influenced artistic practices during the Song Dynasty.")
    
    doc.add_paragraph("This project also aims to analyze the works of major Song Dynasty artists such as Fan Kuan, Guo Xi, and Ma Yuan, and to understand their painting techniques, artistic styles, and cultural meanings. By examining selected artworks and historical context, this study seeks to show how Song Dynasty painting represents harmony between humans and nature.")
    
    doc.add_paragraph("Finally, this project helps develop a deeper appreciation of traditional Chinese art and its lasting influence on later artistic traditions.")
    
    p = doc.add_paragraph("Thank you Jei Wei")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

    # --- Page 3: Introduction ---
    doc.add_heading("The Song Dynasty: The Peak of Chinese Landscape Painting", level=1)
    
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("The Song Dynasty (960-1279 CE) occupies a unique and highly respected place in the history of Chinese art. Unlike earlier periods that emphasized grandeur, power, and court life, Song art focused on nature, inner reflection, and harmony between humans and the universe. Painting during this dynasty reached an extraordinary level of sophistication, especially in landscape painting (shan shui), which became the highest and most respected artistic genre.")
    
    doc.add_paragraph("This project explores the Song Dynasty as a major artistic period, examining its historical background, philosophical foundations, and the works of influential painters such as Fan Kuan, Guo Xi, Li Cheng, Ma Yuan, and Xia Gui. Their artworks reflect the intellectual depth, poetic spirit, and refined aesthetics of Song China.")

    doc.add_heading("Historical and Cultural Background of the Song Dynasty", level=2)
    doc.add_paragraph("The Song Dynasty is divided into two periods: the Northern Song (960-1127) and the Southern Song (1127-1279). During this time, China experienced major advances in education, science, printing technology, and philosophy. Scholar-officials became the dominant social class, and artistic creation was closely connected to scholarship and moral cultivation.")
    
    doc.add_paragraph("Neo-Confucianism, combined with Daoist and Buddhist ideas, strongly influenced Song artists. Painters no longer aimed only to reproduce visible reality; instead, they sought to express inner meaning, emotional atmosphere, and the underlying principles of nature. Painting was regarded as a form of self-cultivation and intellectual reflection rather than simple decoration.")

    doc.add_heading("The Rise of Landscape Painting (Shan Shui)", level=2)
    doc.add_paragraph("Landscape painting reached its highest development during the Song Dynasty. Mountains, rivers, mist, forests, and empty space were used symbolically to express the vastness of the universe and the small yet meaningful position of humans within it.")
    
    p = doc.add_paragraph("Song landscape paintings often emphasize:")
    p.style = 'List Bullet'
    doc.add_paragraph("Monumental mountains and deep spatial depth", style='List Bullet')
    doc.add_paragraph("Subtle use of ink tones rather than bright colors", style='List Bullet')
    doc.add_paragraph("A balance between emptiness and solid forms", style='List Bullet')
    doc.add_paragraph("Strong poetic and philosophical meaning", style='List Bullet')
    
    doc.add_paragraph("These works invite viewers to observe nature quietly and reflect inwardly, encouraging a meditative experience.")

    doc.add_page_break()

    # --- Page 4: Fan Kuan ---
    doc.add_heading("Fan Kuan (c. 960-1030)", level=2)
    doc.add_paragraph("Fan Kuan was one of the most important landscape painters of the Northern Song period. He believed that artists should learn directly from nature instead of copying older artistic styles.")
    
    doc.add_heading("Major Work: Travelers among Mountains and Streams", level=3)
    
    # Insert Fan Kuan Image
    fan_kuan_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/e/e4/Fan_Kuan_-_Travelers_Among_Mountains_and_Streams_-_Google_Art_Project.jpg/512px-Fan_Kuan_-_Travelers_Among_Mountains_and_Streams_-_Google_Art_Project.jpg"
    add_image_from_url(doc, fan_kuan_url, "Fan Kuan - Travelers among Mountains and Streams", width=4.0)

    doc.add_paragraph("This monumental hanging scroll depicts tiny human figures traveling beneath towering mountains. The overwhelming scale of nature compared to humans reflects Daoist ideas of humility and harmony with the natural world. Fan Kuan's strong brushstrokes and textured ink techniques create a sense of realism, strength, and grandeur.")

    # --- Page 5: Guo Xi ---
    doc.add_heading("Guo Xi (c. 1020-1090)", level=2)
    doc.add_paragraph("Guo Xi was both a painter and an important art theorist who served at the imperial court. He believed landscapes should change according to seasons, weather, and viewing perspectives.")
    
    doc.add_heading("Major Work: Early Spring", level=3)

    # Insert Guo Xi Image
    guo_xi_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/5/52/Guo_Xi_-_Early_Spring_%28large%29.jpg/512px-Guo_Xi_-_Early_Spring_%28large%29.jpg"
    add_image_from_url(doc, guo_xi_url, "Guo Xi - Early Spring", width=4.0)
    
    doc.add_paragraph("In Early Spring, Guo Xi presents a landscape filled with mist, rising peaks, and flowing movement. The painting demonstrates his theory of multiple perspectives, allowing the viewer's eye to travel naturally through the scene. The work expresses renewal, balance, and the rhythm of nature.")

    # --- Page 6: Li Cheng ---
    doc.add_heading("Li Cheng (c. 919-967)", level=2)
    doc.add_paragraph("Li Cheng was one of the earliest masters of Northern Song landscape painting and had a major influence on later painters such as Fan Kuan and Guo Xi.")
    
    doc.add_heading("Major Work: A Solitary Temple Amid Clearing Peaks", level=3)

    # Insert Li Cheng Image
    li_cheng_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/6/62/Li_Cheng_-_A_Solitary_Temple_Amid_Clearing_Peaks_-_Google_Art_Project.jpg/512px-Li_Cheng_-_A_Solitary_Temple_Amid_Clearing_Peaks_-_Google_Art_Project.jpg"
    add_image_from_url(doc, li_cheng_url, "Li Cheng - A Solitary Temple Amid Clearing Peaks", width=3.5)
    
    doc.add_paragraph("This painting shows mist-covered mountains with a small temple hidden among them. Li Cheng's soft ink tones and delicate brushwork create a feeling of quiet distance and spiritual calm. His emphasis on space and atmospheric depth reflects Daoist ideals of natural harmony.")

    # --- Page 7: Ma Yuan ---
    doc.add_heading("Ma Yuan (c. 1160-1225)", level=2)
    doc.add_paragraph("Ma Yuan was a leading painter of the Southern Song Dynasty. His style differs greatly from the monumental landscapes of the Northern Song.")
    
    doc.add_heading("Major Work: Walking on a Mountain Path in Spring", level=3)

    # Insert Ma Yuan Image
    ma_yuan_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/c/c5/Ma_Yuan_-_Walking_on_a_Mountain_Path_in_Spring.jpg/640px-Ma_Yuan_-_Walking_on_a_Mountain_Path_in_Spring.jpg"
    add_image_from_url(doc, ma_yuan_url, "Ma Yuan - Walking on a Mountain Path in Spring", width=5.0)

    doc.add_paragraph("Ma Yuan is known for the 'one-corner composition', where most of the visual elements are placed in one area, leaving large empty spaces. This approach creates a poetic and peaceful mood. His paintings focus on emotion, suggestion, and personal experience rather than grand scale.")

    # --- Page 8: Xia Gui ---
    doc.add_heading("Xia Gui (c. 1180-1230)", level=2)
    doc.add_paragraph("Xia Gui was another major painter of the Southern Song Dynasty and a contemporary of Ma Yuan. His style is more energetic and dramatic.")
    
    doc.add_heading("Major Work: Pure and Remote View of Streams and Mountains", level=3)

    # Insert Xia Gui Image
    xia_gui_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/0/06/Pure_and_Remote_View_of_Streams_and_Mountains.jpg/640px-Pure_and_Remote_View_of_Streams_and_Mountains.jpg"
    add_image_from_url(doc, xia_gui_url, "Xia Gui - Pure and Remote View of Streams and Mountains", width=6.0)

    doc.add_paragraph("This handscroll features bold brushstrokes, sharp contrasts, and flowing movement. Xia Gui's use of asymmetrical composition guides the viewer through the landscape dynamically. His work reflects emotional depth and artistic freedom.")

    # Characteristics
    doc.add_heading("Artistic Characteristics of Song Dynasty Painting", level=2)
    doc.add_paragraph("Song Dynasty painting is characterized by:")
    doc.add_paragraph("Dominance of landscape painting as the highest art form", style='List Bullet')
    doc.add_paragraph("Masterful use of ink and brush techniques", style='List Bullet')
    doc.add_paragraph("Philosophical depth influenced by Neo-Confucianism, Daoism, and Buddhism", style='List Bullet')
    doc.add_paragraph("Balance between realism and poetic expression", style='List Bullet')
    doc.add_paragraph("Nature viewed as a moral and spiritual teacher", style='List Bullet')
    
    doc.add_paragraph("These qualities strongly influenced later Chinese and East Asian art traditions.")

    # --- Page 9: Conclusion ---
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("The Song Dynasty represents the intellectual and artistic peak of traditional Chinese painting. Artists such as Fan Kuan, Guo Xi, Li Cheng, Ma Yuan, and Xia Gui transformed landscape painting into a profound visual philosophy. Their works do not simply depict natural scenery; they express ideas about life, harmony, and humanity's place within the universe.")
    
    doc.add_paragraph("Studying Song Dynasty art allows us to understand how painting became a form of meditation, scholarship, and self-expression in Chinese culture.")

    # Save
    doc.save("ArtAppreciation.docx")
    print("Document created successfully as 'ArtAppreciation.docx'")

if __name__ == "__main__":
    create_document()